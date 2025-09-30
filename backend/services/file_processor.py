import base64
import io
import re
from typing import List, Dict, Any

# Try to import optional dependencies
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class FileProcessor:
    """Service for processing different file types and extracting text content"""
    
    @staticmethod
    def process_files(files_data: List[str]) -> List[Dict[str, Any]]:
        """Process a list of files and return processed content for Gemini"""
        processed_files = []
        
        for file_data in files_data:
            try:
                # Parse data URL to get file type and content
                if file_data.startswith('data:'):
                    header, encoded_data = file_data.split(',', 1)
                    mime_type = header.split(';')[0].replace('data:', '')
                    file_content = base64.b64decode(encoded_data)
                    
                    if mime_type.startswith('image/'):
                        # Handle images - pass through to Gemini as is
                        processed_files.append({
                            "mime_type": mime_type,
                            "data": encoded_data
                        })
                    elif mime_type == 'application/pdf':
                        # Extract text from PDF
                        text_content = FileProcessor.extract_pdf_text(file_content)
                        if text_content:
                            processed_files.append({
                                "type": "text",
                                "content": text_content
                            })
                    elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                        # Extract text from DOCX
                        text_content = FileProcessor.extract_docx_text(file_content)
                        if text_content:
                            processed_files.append({
                                "type": "text", 
                                "content": text_content
                            })
                    elif mime_type == 'text/plain':
                        # Handle plain text
                        text_content = file_content.decode('utf-8', errors='ignore')
                        processed_files.append({
                            "type": "text",
                            "content": text_content
                        })
                    else:
                        # Try to decode as text for unknown types
                        try:
                            text_content = file_content.decode('utf-8', errors='ignore')
                            processed_files.append({
                                "type": "text",
                                "content": text_content
                            })
                        except:
                            continue
                            
            except Exception as e:
                # Skip files that can't be processed
                print(f"Error processing file: {e}")
                continue
        
        return processed_files
    
    @staticmethod
    def extract_pdf_text(pdf_content: bytes) -> str:
        """Extract text from PDF content"""
        if not PDF_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise ImportError("PDF processing libraries not available. Please install PyPDF2 or pdfplumber.")
        
        text_content = ""
        
        # Try pdfplumber first (better text extraction)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
                return text_content.strip()
            except Exception as e:
                print(f"pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n\n"
                return text_content.strip()
            except Exception as e:
                print(f"PyPDF2 failed: {e}")
        
        return ""
    
    @staticmethod
    def extract_docx_text(docx_content: bytes) -> str:
        """Extract text from DOCX content"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(docx_content))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""
    
    @staticmethod
    def create_mixed_content_for_gemini(processed_files: List[Dict[str, Any]]) -> List[Any]:
        """Create content array for Gemini API with mixed text and images"""
        content_parts = []
        
        # Combine all text content
        text_parts = []
        for file_data in processed_files:
            if file_data.get("type") == "text":
                text_parts.append(file_data["content"])
        
        # Add combined text if any
        if text_parts:
            combined_text = "\n\n--- Document Separator ---\n\n".join(text_parts)
            content_parts.append(f"Text content from uploaded documents:\n\n{combined_text}")
        
        # Add images
        for file_data in processed_files:
            if "mime_type" in file_data:  # This is an image
                content_parts.append({
                    "mime_type": file_data["mime_type"],
                    "data": file_data["data"]
                })
        
        return content_parts
