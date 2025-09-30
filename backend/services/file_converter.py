import re
from io import BytesIO
from typing import Dict

# Try to import optional dependencies
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

class FileConverter:
    """Service for converting MOM content to different file formats"""
    
    @staticmethod
    def markdown_to_txt(content: str) -> str:
        """Convert markdown content to plain text"""
        # Remove markdown formatting
        text = content
        
        # Remove markdown headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove markdown bold/italic
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Convert markdown lists to plain text
        text = re.sub(r'^[\s]*[-*+]\s+', '• ', text, flags=re.MULTILINE)
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)
        
        # Remove markdown links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Convert markdown tables to plain text
        lines = text.split('\n')
        converted_lines = []
        in_table = False
        
        for line in lines:
            # Check if line is a table row
            if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
                if not in_table:
                    in_table = True
                # Convert table row to plain text
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if not all(cell == '---' or cell == '' for cell in cells):  # Skip separator rows
                    converted_lines.append(' | '.join(cells))
            else:
                if in_table:
                    in_table = False
                    converted_lines.append('')  # Add blank line after table
                converted_lines.append(line)
        
        text = '\n'.join(converted_lines)
        
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = text.strip()
        
        return text
    
    @staticmethod
    def markdown_to_docx(content: str) -> BytesIO:
        """Convert markdown content to DOCX format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Add blank paragraph for empty lines
                doc.add_paragraph()
                i += 1
                continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('# ').strip()
                
                if level == 1:
                    # Main heading
                    heading = doc.add_heading(header_text, level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    # Sub headings
                    doc.add_heading(header_text, level=min(level, 3))
                
            # Handle tables
            elif line.startswith('|') and line.endswith('|'):
                table_lines = []
                j = i
                
                # Collect all table lines
                while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                    table_line = lines[j].strip()
                    # Skip separator lines
                    if not all(cell.strip() in ['---', ''] for cell in table_line.split('|')[1:-1]):
                        table_lines.append(table_line)
                    j += 1
                
                if table_lines:
                    # Create table
                    first_row = table_lines[0]
                    cols = len([cell.strip() for cell in first_row.split('|')[1:-1]])
                    rows = len(table_lines)
                    
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, table_line in enumerate(table_lines):
                        cells_data = [cell.strip() for cell in table_line.split('|')[1:-1]]
                        for col_idx, cell_data in enumerate(cells_data):
                            if col_idx < cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_data
                                # Make header row bold
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                
                i = j
                continue
            
            # Handle lists
            elif line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.\s', line):
                # Bullet or numbered list
                if line.startswith(('- ', '* ', '+ ')):
                    list_text = line[2:].strip()
                    p = doc.add_paragraph(list_text, style='List Bullet')
                else:
                    list_text = re.sub(r'^\d+\.\s+', '', line)
                    p = doc.add_paragraph(list_text, style='List Number')
            
            # Handle bold text formatting
            else:
                # Regular paragraph with formatting
                p = doc.add_paragraph()
                
                # Process bold and other formatting
                text_parts = re.split(r'(\*\*[^*]+\*\*)', line)
                
                for part in text_parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        bold_text = part[2:-2]
                        run = p.add_run(bold_text)
                        run.bold = True
                    else:
                        # Regular text
                        p.add_run(part)
            
            i += 1
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
